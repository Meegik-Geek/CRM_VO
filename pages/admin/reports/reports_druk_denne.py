
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH

from tkinter import Tk, filedialog
from datetime import datetime
from db.connect_db import setup_database, close_database, get_setting
from openpyxl import Workbook
class DocumentPrinter:
    def __init__(self, show_success_message, show_error_message):
        """Ініціалізація принтера з повідомленнями про успіх і помилки."""
        self.conn = None
        self.cursor = None
        self.show_success_message = show_success_message
        self.show_error_message = show_error_message

    def connect_db(self):
        """Відкриває підключення до бази даних."""
        self.conn = setup_database()
        self.cursor = self.conn.cursor()

    def close_db(self):
        """Закриває підключення до бази даних."""
        if self.conn:
            close_database(self.conn)

    def _get_common_data(self):
        """Повертає загальні дані для всіх звітів."""
        return {
            "institution_name": get_setting("college_name", "Назва закладу"),
            "institution_short_name": get_setting("college_short_name", "Скорочена назва"),
            "resp_secretary": get_setting("resp_secretary", "Відповідальний секретар"),
            "deputy_secretary": get_setting("deputy_secretary", "Заступник відповідального секретаря"),
            "legal_counsel": get_setting("legal_counsel", "Юрисконсульт"),
            "edebo_admin": get_setting("edebo_admin", "Відповідальний за ЄДЕБО"),
        }

    def fetch_data(self, query, params ):
        """Виконує SQL-запит і повертає дані."""
        self.cursor.execute(query, params)
        return self.cursor.fetchone()

    def _apply_chunky_sorting(self, applicants, group_size):
        """Допоміжний метод: сортує абітурієнтів за алфавітом всередині кожної групи."""
        if group_size and group_size > 0:
            final_sorted_list = []
            for i in range(0, len(applicants), group_size):
                chunk = list(applicants[i:i + group_size])
                # Сортуємо пачку за ПІБ (перший елемент кортежу)
                chunk.sort(key=lambda x: x[0])
                final_sorted_list.extend(chunk)
            return final_sorted_list
        return applicants

    def _set_row_color(self, row, hex_color):
        """Встановлює колір заливки для всіх комірок у рядку таблиці Word."""
        for cell in row.cells:
            tcPr = cell._element.get_or_add_tcPr()
            # Створюємо XML для заливки (shading)
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
            tcPr.append(shd)

    def print_dopusk_page(self, specialty_name, group_size, dialog):
        """Друк звіту 'Допуск до вступних випробувань' з динамічним розбиттям на групи та підгрупи."""
        try:
            self.connect_db()
            query = """
                SELECT 
                    pd.last_name || ' ' || pd.first_name || ' ' || COALESCE(pd.middle_name, '') AS full_name,
                    os.number_sprava
                FROM 
                    public.applicant_personal_data_day pd
                JOIN 
                    public.personal_case_day os ON pd.cert_number = os.cert_number
                WHERE 
                    os.name_specialnosti = %s
                    AND (os.is_cancelled IS NOT TRUE)
                ORDER BY 
                    -- Надійне сортування: спочатку за довжиною (щоб 2 < 10), потім за значенням
                    LENGTH(os.number_sprava) ASC,
                    os.number_sprava ASC,
                    pd.last_name ASC
            """
            
            self.cursor.execute(query, (specialty_name,))
            applicants = self.cursor.fetchall()
            
            if not applicants:
                self.show_error_message(f"Не знайдено абітурієнтів для спеціальності '{specialty_name}'.")
                self.close_db()
                return
            
            # Застосовуємо алгоритм групового алфавітного сортування
            applicants = self._apply_chunky_sorting(applicants, group_size)

            # Формування динамічних груп
            groups_context = []
            for i in range(0, len(applicants), group_size):
                group_raw = applicants[i:i + group_size]
                # Сортуємо в межах групи за прізвищем (якщо потрібно)
                # group_raw.sort(key=lambda x: x[0].split()[0])
                
                sub1_size = (len(group_raw) + 1) // 2
                sub1 = group_raw[:sub1_size]
                sub2 = group_raw[sub1_size:]
                
                groups_context.append({
                    "number": f"{(i // group_size) + 1:02d}",
                    "sub1": [{"index": j + 1, "full_name": s[0]} for j, s in enumerate(sub1)],
                    "sub2": [{"index": sub1_size + j + 1, "full_name": s[0]} for j, s in enumerate(sub2)]
                })

            template_path = "templates/dopusk_template.docx"
            context = {
                "specialty_name": specialty_name,
                "groups": groups_context
            }

            self.fill_and_save_template(template_path, context, dialog, "Документ успішно збережено!")
            self.close_db()

        except Exception as e:
            self.show_error_message(f"Помилка при створенні документів: {str(e)}")
            self.close_db()
  
       

    def print_dopusk_zaoch_page(self, specialty_name, group_size, dialog):
        """Друк звіту 'Допуск до вступних випробувань' для заочної форми з динамічними групами."""
        try:
            self.connect_db()
            query = """
                SELECT 
                    pd.last_name || ' ' || pd.first_name || ' ' || COALESCE(pd.middle_name, '') AS full_name,
                    os.date_sprava
                FROM 
                    public.applicant_personal_data_evening pd
                JOIN 
                    public.personal_case_evening os ON pd.cert_number = os.cert_number
                WHERE 
                    os.name_specialnosti = %s
                    AND (os.zno_nmt_checkbox IS NULL OR os.zno_nmt_checkbox != 'true')
                    AND (os.is_cancelled IS NOT TRUE)
                ORDER BY 
                    LENGTH(os.number_sprava) ASC,
                    os.number_sprava ASC,
                    pd.last_name ASC
            """
            
            self.cursor.execute(query, (specialty_name,))
            applicants = self.cursor.fetchall()
            
            if not applicants:
                self.show_error_message(f"Немає абітурієнтів без сертифіката НМТ для '{specialty_name}' (заочна).")
                self.close_db()
                return

            # Застосовуємо алгоритм групового алфавітного сортування
            applicants = self._apply_chunky_sorting(applicants, group_size)

            # Формування динамічних груп
            groups_context = []
            for i in range(0, len(applicants), group_size):
                group_raw = applicants[i:i + group_size]
                
                sub1_size = (len(group_raw) + 1) // 2
                sub1 = group_raw[:sub1_size]
                sub2 = group_raw[sub1_size:]
                
                groups_context.append({
                    "number": f"{(i // group_size) + 1:02d}",
                    "sub1": [{"index": j + 1, "full_name": s[0]} for j, s in enumerate(sub1)],
                    "sub2": [{"index": sub1_size + j + 1, "full_name": s[0]} for j, s in enumerate(sub2)]
                })

            template_path = "templates/dopusk_template.docx"
            context = {
                "specialty_name": specialty_name + " (заочна)",
                "groups": groups_context
            }

            self.fill_and_save_template(template_path, context, dialog, "Документ успішно збережено!")
            self.close_db()

        except Exception as e:
            self.show_error_message(f"Помилка при створенні документів: {str(e)}")
            self.close_db()

    def print_dopusk_scor_page(self, specialty_name, group_size, dialog):
        """Друк звіту 'Допуск до вступних випробувань' для денної скороченої форми з динамічними групами."""
        try:
            self.connect_db()
            query = """
                SELECT 
                    pd.last_name || ' ' || pd.first_name || ' ' || COALESCE(pd.middle_name, '') AS full_name,
                    os.date_sprava
                FROM 
                    public.applicant_personal_data_day pd
                JOIN 
                    public.personal_case_day_scor os ON pd.cert_number = os.cert_number
                WHERE 
                    os.name_specialnosti = %s
                    AND (os.zno_nmt_checkbox IS NULL OR os.zno_nmt_checkbox != 'true')
                    AND (os.is_cancelled IS NOT TRUE)
                ORDER BY 
                    LENGTH(os.number_sprava) ASC,
                    os.number_sprava ASC,
                    pd.last_name ASC
            """
            
            self.cursor.execute(query, (specialty_name,))
            applicants = self.cursor.fetchall()
            
            if not applicants:
                self.show_error_message(f"Немає абітурієнтів без сертифіката НМТ для '{specialty_name}' (скорочена денна).")
                self.close_db()
                return

            # Застосовуємо алгоритм групового алфавітного сортування
            applicants = self._apply_chunky_sorting(applicants, group_size)

            # Формування динамічних груп
            groups_context = []
            for i in range(0, len(applicants), group_size):
                group_raw = applicants[i:i + group_size]
                
                sub1_size = (len(group_raw) + 1) // 2
                sub1 = group_raw[:sub1_size]
                sub2 = group_raw[sub1_size:]
                
                groups_context.append({
                    "number": f"{(i // group_size) + 1:02d}",
                    "sub1": [{"index": j + 1, "full_name": s[0]} for j, s in enumerate(sub1)],
                    "sub2": [{"index": sub1_size + j + 1, "full_name": s[0]} for j, s in enumerate(sub2)]
                })

            template_path = "templates/dopusk_template.docx"
            context = {
                "specialty_name": specialty_name + " (денна)",
                "groups": groups_context
            }

            self.fill_and_save_template(template_path, context, dialog, "Документ успішно збережено!")
            self.close_db()

        except Exception as e:
            self.show_error_message(f"Помилка при створенні документів: {str(e)}")
            self.close_db()

    def print_results_exams_all_forms(self, specialty_name, group_size, form_type, dialog):
        """Друк результатів вступних випробувань для різних форм навчання."""
        try:
            self.connect_db()
            
            # Визначаємо таблиці в залежності від форми навчання
            if form_type == 'day':
                table_cases = "public.personal_case_day"
                table_personal = "public.applicant_personal_data_day"
                spec_display = specialty_name + " (денна)"
            elif form_type == 'day_scor':
                table_cases = "public.personal_case_day_scor"
                table_personal = "public.applicant_personal_data_day"
                spec_display = specialty_name + " (денна скорочена)"
            elif form_type == 'evening':
                table_cases = "public.personal_case_evening"
                table_personal = "public.applicant_personal_data_evening"
                spec_display = specialty_name + " (заочна)"
            else:
                self.show_error_message("Невідома форма навчання.")
                self.close_db()
                return

            # Сортування: ЗАВЖДИ як у Допуску
            order_by = "LENGTH(os.number_sprava) ASC, os.number_sprava ASC"

            query = f"""
                SELECT 
                    pd.last_name || ' ' || pd.first_name || ' ' || COALESCE(pd.middle_name, '') AS full_name,
                    os.number_sprava,
                    es.score,
                    es.status
                FROM 
                    {table_cases} os
                JOIN 
                    {table_personal} pd ON os.cert_number = pd.cert_number
                LEFT JOIN 
                    entrance_scores es ON TRIM(os.number_sprava) = TRIM(es.number_sprava)
                WHERE 
                    os.name_specialnosti = %s
                    AND (os.is_cancelled IS NOT TRUE)
                ORDER BY 
                    {order_by}, pd.last_name ASC
            """
            
            self.cursor.execute(query, (specialty_name,))
            applicants = self.cursor.fetchall()
            
            if not applicants:
                self.show_error_message(f"Не знайдено абітурієнтів для спеціальності '{specialty_name}' ({form_type}).")
                self.close_db()
                return
            
            # Застосовуємо алгоритм групового алфавітного сортування
            applicants = self._apply_chunky_sorting(applicants, group_size)

            # Підготовка контексту для шаблону: ЗАВЖДИ плоский список
            items = []
            for i, s in enumerate(applicants):
                score_val = s[2]
                status_val = s[3] if s[3] else ""
                
                # Логіка: якщо балів 0 або None (немає запису), показуємо Статус
                display_score = str(score_val) if (score_val is not None and score_val != 0) else status_val
                
                items.append({
                    "index": i + 1,
                    "full_name": s[0],
                    "score": display_score
                })
            
            context = {
                **self._get_common_data(),
                "specialty_name": spec_display,
                "items": items
            }

            template_path = "templates/rezult_vstup_viprobov.docx"
            self.fill_and_save_template(template_path, context, dialog, "Результати випробувань успішно збережено!")
            self.close_db()

        except Exception as e:
            self.show_error_message(f"Помилка при створенні звіту: {str(e)}")
            self.close_db()

    def print_motivation_ranking_all_forms(self, specialty_name, form_type, dialog):
        """Друк ранжування мотиваційних листів для різних форм навчання."""
        try:
            self.connect_db()
            
            # Визначаємо таблиці
            if form_type == 'day':
                table_cases = "public.personal_case_day"
                table_personal = "public.applicant_personal_data_day"
                spec_display = specialty_name + " (денна)"
            elif form_type == 'day_scor':
                table_cases = "public.personal_case_day_scor"
                table_personal = "public.applicant_personal_data_day"
                spec_display = specialty_name + " (денна скорочена)"
            elif form_type == 'evening':
                table_cases = "public.personal_case_evening"
                table_personal = "public.applicant_personal_data_evening"
                spec_display = specialty_name + " (заочна)"
            else:
                self.show_error_message("Невідома форма навчання.")
                self.close_db()
                return

            # Сортування: за Балом (DESC) та Номером справи (ASC)
            order_by = "es.score DESC, LENGTH(os.number_sprava) ASC, os.number_sprava ASC"

            query = f"""
                SELECT 
                    pd.last_name || ' ' || pd.first_name || ' ' || COALESCE(pd.middle_name, '') AS full_name,
                    os.number_sprava,
                    es.score,
                    es.motivation_rank
                FROM 
                    {table_cases} os
                JOIN 
                    {table_personal} pd ON os.cert_number = pd.cert_number
                LEFT JOIN 
                    entrance_scores es ON TRIM(os.number_sprava) = TRIM(es.number_sprava)
                WHERE 
                    os.name_specialnosti = %s
                    AND (os.is_cancelled IS NOT TRUE)
                ORDER BY 
                    {order_by}
            """
            
            self.cursor.execute(query, (specialty_name,))
            applicants = self.cursor.fetchall()
            
            if not applicants:
                self.show_error_message(f"Не знайдено абітурієнтів для спеціальності '{specialty_name}' ({form_type}).")
                self.close_db()
                return
            
            # Формуємо список даних з фільтрацією (тільки ті, у кого однакові бали)
            from collections import Counter
            all_scores = [s[2] for s in applicants if s[2] is not None]
            score_counts = Counter(all_scores)
            
            # Залишаємо тільки тих, чий бал зустрічається 2 або більше разів
            filtered_applicants = [s for s in applicants if s[2] is not None and score_counts[s[2]] > 1]
            
            if not filtered_applicants:
                self.show_error_message(f"Не знайдено абітурієнтів з однаковими балами для спеціальності '{specialty_name}'.")
                self.close_db()
                return

            items = []
            for i, s in enumerate(filtered_applicants):
                items.append({
                    "index": i + 1,
                    "full_name": s[0],
                    "score": s[2],
                    "motivation_rank": s[3] if s[3] is not None else ""
                })
            
            context = {
                **self._get_common_data(),
                "specialty_name": spec_display,
                "items": items
            }

            template_path = "templates/rezult_motivation_ranking.docx"
            self.fill_and_save_template(template_path, context, dialog, "Ранжування успішно збережено!")
            self.close_db()

        except Exception as e:
            self.show_error_message(f"Помилка при створенні ранжування: {str(e)}")
            self.close_db()

    def print_rating_list_all_forms(self, specialty_name, form_type, dialog):
        """Друк рейтингових списків з пріоритетами та відсічками."""
        try:
            self.connect_db()
            
            # 1. Отримуємо ліміти спеціальності
            if form_type == 'evening':
                table_spec = "public.specialities_evening"
            else:
                table_spec = "public.specialities_day"
                
            self.cursor.execute(f"SELECT state_places, licensed_volume FROM {table_spec} WHERE name_specialnosti = %s", (specialty_name,))
            limits = self.cursor.fetchone()
            state_places = limits[0] if limits else 0
            licensed_volume = limits[1] if limits else 0

            # 2. Визначаємо таблиці для вступників
            if form_type == 'day':
                table_cases = "public.personal_case_day"
                table_personal = "public.applicant_personal_data_day"
                table_benefits = "public.applicant_benefits_day"
                spec_display = specialty_name + " (денна)"
            elif form_type == 'day_scor':
                table_cases = "public.personal_case_day_scor"
                table_personal = "public.applicant_personal_data_day"
                table_benefits = "public.applicant_benefits_day" # Використовуємо денну пільгу
                spec_display = specialty_name + " (денна скорочена)"
            elif form_type == 'evening':
                table_cases = "public.personal_case_evening"
                table_personal = "public.applicant_personal_data_evening"
                table_benefits = "public.applicant_benefits_evening"
                spec_display = specialty_name + " (заочна)"
            else:
                self.show_error_message("Невідома форма навчання.")
                self.close_db()
                return

            # 3. SQL з пріоритетним сортуванням
            query = f"""
                WITH interview_data AS (
                    SELECT ab.cert_number, 
                           bool_or(b.type_pilgi = 'Вступ за співбесідою') as is_interview
                    FROM {table_benefits} ab
                    JOIN benefits b ON ab.kod_pilgi = b.kod_pilgi
                    GROUP BY ab.cert_number
                )
                SELECT 
                    pd.last_name || ' ' || pd.first_name || ' ' || COALESCE(pd.middle_name, '') AS full_name,
                    os.number_sprava,
                    es.score,
                    es.motivation_rank,
                    COALESCE(id.is_interview, false) as is_interview
                FROM 
                    {table_cases} os
                JOIN 
                    {table_personal} pd ON os.cert_number = pd.cert_number
                LEFT JOIN 
                    entrance_scores es ON TRIM(os.number_sprava) = TRIM(es.number_sprava)
                LEFT JOIN 
                    interview_data id ON os.cert_number = id.cert_number
                WHERE 
                    os.name_specialnosti = %s
                    AND (os.is_cancelled IS NOT TRUE)
                ORDER BY 
                    CASE WHEN COALESCE(id.is_interview, false) THEN 0 ELSE 1 END ASC,
                    es.score DESC NULLS LAST,
                    es.motivation_rank ASC NULLS LAST,
                    LENGTH(os.number_sprava) ASC,
                    os.number_sprava ASC
            """
            
            self.cursor.execute(query, (specialty_name,))
            applicants = self.cursor.fetchall()
            
            if not applicants:
                self.show_error_message(f"Не знайдено абітурієнтів для спеціальності '{specialty_name}' ({form_type}).")
                self.close_db()
                return
            
            # 4. Формуємо список та конфігурацію кольорів
            items = []
            colors_config = {} # {номер_рядка: hex_колір}
            
            for i, s in enumerate(applicants):
                real_index = i + 1
                
                # Підготовка відображення бала
                if s[4]: # is_interview
                    display_score = "за результатами співбесіди"
                else:
                    display_score = str(s[2]) if s[2] is not None else ""

                items.append({
                    "index": real_index,
                    "full_name": s[0],
                    "display_score": display_score
                })

                # Визначаємо колір (рядки в таблиці Word починаються з 0, дані з 1)
                row_idx = real_index
                if real_index <= state_places:
                    colors_config[row_idx] = "92FCA6"  # Яскраво-зелений (Бюджет)
                elif real_index <= licensed_volume:
                    colors_config[row_idx] = "B4D4DA"  # Блакитний (Контракт)
            
            context = {
                **self._get_common_data(),
                "specialty_name": spec_display,
                "date": datetime.now().strftime("%d.%m.%Y"),
                "items": items
            }

            template_path = "templates/rating_list_template.docx"
            self.fill_and_save_template(template_path, context, dialog, "Рейтинговий список успішно збережено!", colors_config=colors_config)
            self.close_db()

        except Exception as e:
            self.show_error_message(f"Помилка при створенні рейтингу: {str(e)}")
            self.close_db()
    def print_registr_vstupnik_page(self):
        """Друк загального журналу реєстрації вступників по всіх спеціальностях."""
        try:
            self.connect_db()
            
            # SQL-запит для отримання інформації про вступників
            query = """
            SELECT * FROM (
                SELECT 
                    pd.last_name || ' ' || pd.first_name || ' ' || COALESCE(pd.middle_name, '') AS full_name,
                    os.date_sprava AS date_sprava,  -- Повна дата подачі
                    os.number_sprava,  -- Номер справи
                    os.name_specialnosti,  -- Назва спеціальності
                    pd.cert_number  -- Свідоцтво про освіту
                FROM 
                    public.applicant_personal_data_day pd
                JOIN 
                    public.personal_case_day os 
                ON 
                    pd.cert_number = os.cert_number

                UNION ALL

                SELECT 
                    pd.last_name || ' ' || pd.first_name || ' ' || COALESCE(pd.middle_name, '') AS full_name,
                    os.date_sprava AS date_sprava,  -- Повна дата подачі
                    os.number_sprava,  -- Номер справи
                    os.name_specialnosti,  -- Назва спеціальності
                    pd.cert_number  -- Свідоцтво про освіту
                FROM 
                    public.applicant_personal_data_day pd
                JOIN 
                    public.personal_case_day_scor os 
                ON 
                    pd.cert_number = os.cert_number

                UNION ALL

                SELECT 
                    pd.last_name || ' ' || pd.first_name || ' ' || COALESCE(pd.middle_name, '') AS full_name,
                    os.date_sprava AS date_sprava,  -- Повна дата подачі
                    os.number_sprava,  -- Номер справи
                    os.name_specialnosti,  -- Назва спеціальності
                    pd.cert_number  -- Свідоцтво про освіту
                FROM 
                    public.applicant_personal_data_evening pd
                JOIN 
                    public.personal_case_evening os 
                ON 
                    pd.cert_number = os.cert_number
            ) AS combined_data
            ORDER BY 
                date_sprava ASC, 
                full_name ASC;
            """
            
            self.cursor.execute(query)
            raw_applicants = self.cursor.fetchall()

            if not raw_applicants:
                self.show_error_message("Немає даних для журналу реєстрації вступників.")
                self.close_db()
                return

            # Організовуємо дані на рівні Python
            applicants_dict = {}
            for full_name, date_sprava, number_sprava, specialty_name, cert_number in raw_applicants:
                # Перевіряємо і додаємо вступника в словник
                if cert_number not in applicants_dict:
                    applicants_dict[cert_number] = {
                        "full_name": full_name,
                        "first_date_sprava": date_sprava,
                        "cases": [],  # Список для пар (номер справи, спеціальність)
                        "cert_number": cert_number
                    }
                applicants_dict[cert_number]["cases"].append((number_sprava, specialty_name))

            # Формуємо список заявників для шаблону
            context = {"applicants": []}
            for i, (cert_number, applicant_data) in enumerate(applicants_dict.items(), start=1):
                # Форматуємо дату подачі
                date_sprava = applicant_data["first_date_sprava"]
                if date_sprava is not None:
                    date_sprava_str = (
                        date_sprava
                        if isinstance(date_sprava, str)
                        else date_sprava.strftime("%d.%m.%Y")
                    )
                else:
                    date_sprava_str = "Н/Д"

                # Форматуємо номери справ і спеціальності
                numbers_sprava = ", ".join(set(case[0] for case in applicant_data["cases"]))
                specialties = ", ".join(set(case[1] for case in applicant_data["cases"]))

                # Додаємо у контекст
                context["applicants"].append({
                    "index": i,
                    "full_name": applicant_data["full_name"],
                    "number_sprava": numbers_sprava,
                    "specialty_name": specialties,
                    "date_sprava": date_sprava_str,
                    "cert_number": cert_number
                })

            # Заповнюємо шаблон і зберігаємо документ
            template_path = "templates/registr_vstupnik_template.docx"
            self.fill_and_save_template_dialog(template_path, context, "Журнал реєстрації вступників успішно збережений!")
            self.close_db()

        except Exception as e:
            self.show_error_message(f"Помилка при створенні журналу реєстрації: {str(e)}")
            self.close_db()



    def print_alfavit_vstupnik_page(self):
        """Друк алфавітного журналу вступників по всіх спеціальностях (кожна буква - окрема таблиця з початком нумерації з 1)."""
        try:
            self.connect_db()

            # SQL-запит для отримання інформації про вступників
            query = """
                SELECT 
                    pd.last_name || ' ' || pd.first_name || ' ' || COALESCE(pd.middle_name, '') AS full_name,
                    os.number_sprava,
                    os.name_specialnosti
                FROM 
                    public.applicant_personal_data_day pd 
                JOIN 
                    public.personal_case_day os 
                ON 
                    pd.cert_number = os.cert_number

                UNION ALL

                SELECT 
                    pd.last_name || ' ' || pd.first_name || ' ' || COALESCE(pd.middle_name, '') AS full_name,
                    os.number_sprava,
                    os.name_specialnosti
                FROM 
                    public.applicant_personal_data_day pd 
                JOIN 
                    public.personal_case_day_scor os 
                ON 
                    pd.cert_number = os.cert_number

                UNION ALL

                SELECT 
                    pd.last_name || ' ' || pd.first_name || ' ' || COALESCE(pd.middle_name, '') AS full_name,
                    os.number_sprava,
                    os.name_specialnosti
                FROM 
                    public.applicant_personal_data_evening pd 
                JOIN 
                    public.personal_case_evening os 
                ON 
                    pd.cert_number = os.cert_number

                ORDER BY 
                    full_name ASC;
            """
            
            self.cursor.execute(query)
            raw_applicants = self.cursor.fetchall()

            if not raw_applicants:
                self.show_error_message("Немає даних для алфавітного журналу вступників.")
                self.close_db()
                return

            # Групування вступників за повним ім'ям, зберігаючи відповідні пари "номер справи - спеціальність"
            applicants_dict = {}
            for full_name, number_sprava, specialty_name in raw_applicants:
                if full_name not in applicants_dict:
                    applicants_dict[full_name] = {
                        "full_name": full_name,
                        "cases": []  # Список для пар (номер справи, спеціальність)
                    }
                applicants_dict[full_name]["cases"].append((number_sprava, specialty_name))

            # Формуємо список заявників для вставки у шаблон
            context = {"applicants_by_letter": []}
            applicants_by_letter = {}

            # Заповнюємо дані для шаблону з новою нумерацією для кожної букви
            for full_name, applicant_data in applicants_dict.items():
                # Об'єднуємо відповідні номери справ і спеціальності
                numbers_sprava = ", ".join(case[0] for case in applicant_data["cases"])
                specialties = ", ".join(case[1] for case in applicant_data["cases"])

                # Групуємо за першою літерою прізвища
                first_letter = full_name[0].upper()
                if first_letter not in applicants_by_letter:
                    applicants_by_letter[first_letter] = []

                # Додаємо запис із ПІБ, номерами справ і спеціальностями
                applicant_entry = {
                    "full_name": full_name,
                    "number_sprava": numbers_sprava,
                    "specialty_name": specialties,
                }
                applicants_by_letter[first_letter].append(applicant_entry)

            # Формуємо контекст для шаблону, де кожна нова літера має свій початковий номер
            for letter in sorted(applicants_by_letter.keys()):
                letter_section = {"letter": letter, "applicants": []}
                
                # Оновлюємо індекси, починаючи з 1 для кожної літери
                for i, applicant in enumerate(applicants_by_letter[letter], start=1):
                    applicant["index"] = i
                    letter_section["applicants"].append(applicant)
                
                context["applicants_by_letter"].append(letter_section)

            # Заповнюємо шаблон і зберігаємо документ
            template_path = "templates/alfavit_vstupnik_template.docx"
            self.fill_and_save_template_dialog(template_path, context, "Алфавітний журнал вступників успішно збережений!")
            self.close_db()

        except Exception as e:
            self.show_error_message(f"Помилка при створенні журналу: {str(e)}")
            self.close_db()



    def print_spec_umovi_vstupnik_page(self):
        """Друк журналу реєстрації вступників із спеціальними умовами (пільгами) для всіх спеціальностей."""
        try:
            self.connect_db()

            # SQL-запит для отримання інформації про вступників із пільгами, виключаючи записи без коду пільги
            query = """
                SELECT * FROM (
                    SELECT 
                        pd.last_name || ' ' || pd.first_name || ' ' || COALESCE(pd.middle_name, '') AS full_name,
                        MIN(os.date_sprava) AS first_date_sprava,
                        os.number_sprava,
                        os.name_specialnosti,
                        pd.cert_number,
                        pd.address,
                        pd.gender,
                        pd.date_birth,
                        pv.kod_pilgi,
                        pv.document_pilgi
                    FROM 
                        public.applicant_personal_data_day pd
                    JOIN 
                        public.personal_case_day os ON pd.cert_number = os.cert_number
                    LEFT JOIN 
                        public.applicant_benefits_day pv ON pd.cert_number = pv.cert_number
                    WHERE 
                        pv.kod_pilgi IS NOT NULL
                    GROUP BY 
                         pd.last_name, pd.first_name, pd.middle_name, pd.cert_number, 
                            pd.address, pd.gender, pd.date_birth,
                            os.number_sprava, os.name_specialnosti, 
                            pv.kod_pilgi, pv.document_pilgi

                    UNION ALL

                    SELECT 
                        pd.last_name || ' ' || pd.first_name || ' ' || COALESCE(pd.middle_name, '') AS full_name,
                        MIN(os.date_sprava) AS first_date_sprava,
                        os.number_sprava,
                        os.name_specialnosti,
                        pd.cert_number,
                        pd.address,
                        pd.gender,
                        pd.date_birth,
                        pv.kod_pilgi,
                        pv.document_pilgi
                    FROM 
                        public.applicant_personal_data_evening pd
                    JOIN 
                        public.personal_case_evening os ON pd.cert_number = os.cert_number
                    LEFT JOIN 
                        public.applicant_benefits_evening pv ON pd.cert_number = pv.cert_number
                    WHERE 
                        pv.kod_pilgi IS NOT NULL
                    GROUP BY 
                        pd.last_name, pd.first_name, pd.middle_name, pd.cert_number, 
                            pd.address, pd.gender, pd.date_birth,
                            os.number_sprava, os.name_specialnosti, 
                            pv.kod_pilgi, pv.document_pilgi
                ) AS combined_data
                ORDER BY 
                    first_date_sprava ASC, 
                    full_name ASC;

            """

            self.cursor.execute(query)
            raw_applicants = self.cursor.fetchall()

            # Організовуємо дані на рівні Python, щоб уникнути дублікатів і зберегти відповідність
            applicants_dict = {}
            for row in raw_applicants:
                full_name, first_date_sprava, number_sprava, specialty_name, cert_number, address, gender, date_birth, kod_pilgi, document_pilgi = row

                # Перевірка наявності абітурієнта в словнику
                if cert_number not in applicants_dict:
                    applicants_dict[cert_number] = {
                        "full_name": full_name,
                        "first_date_sprava": first_date_sprava,
                        "cases": [],  # Список для пар (номер справи, спеціальність)
                        "address": address,
                        "gender": gender,
                        "date_birth": date_birth,
                        "cert_number": cert_number,
                        "pilgi_info": [],  # Зберігаємо кілька пар (код пільги, документ пільги)
                    }

                # Додаємо нову справу і спеціальність лише якщо ще не додано
                if (number_sprava, specialty_name) not in applicants_dict[cert_number]["cases"]:
                    applicants_dict[cert_number]["cases"].append((number_sprava, specialty_name))

                # Додаємо код пільги і документ пільги лише якщо ще не додано
                if (kod_pilgi, document_pilgi) not in applicants_dict[cert_number]["pilgi_info"]:
                    applicants_dict[cert_number]["pilgi_info"].append((kod_pilgi, document_pilgi))

            # Формуємо список заявників для вставки у шаблон
            context = {"applicants": []}
            for i, (cert_number, applicant_data) in enumerate(applicants_dict.items(), start=1):
                # Перевірка та форматування first_date_sprava
                date_sprava_str = (
                    applicant_data["first_date_sprava"]
                    if isinstance(applicant_data["first_date_sprava"], str)
                    else applicant_data["first_date_sprava"].strftime("%d.%m.%Y")
                )

                # Перевірка та форматування date_birth
                date_birth_str = (
                    applicant_data["date_birth"]
                    if isinstance(applicant_data["date_birth"], str)
                    else applicant_data["date_birth"].strftime("%d.%m.%Y")
                )

                # Форматуємо номери справ і спеціальності відповідно до їх порядку
                numbers_sprava = ", ".join(case[0] for case in applicant_data["cases"])
                specialties = ", ".join(case[1] for case in applicant_data["cases"])
                pilgi_info_str = "; ".join(f"{code}: {doc}" for code, doc in applicant_data["pilgi_info"])

                # Додаємо кожного абітурієнта у контекст для шаблону
                context["applicants"].append({
                    "index": i,
                    "full_name": applicant_data["full_name"],
                    "number_sprava": numbers_sprava,
                    "specialty_name": specialties,
                    "date_sprava": date_sprava_str,
                    "cert_number": cert_number,
                    "address": applicant_data["address"],
                    "gender_birth": f"{applicant_data['gender']} / {date_birth_str}",
                    "pilgi_info": pilgi_info_str or "Немає інформації",
                    "remark": ", ".join(str(kod_pilgi) for kod_pilgi, _ in applicant_data["pilgi_info"])
                })

            # Заповнюємо шаблон і зберігаємо документ
            template_path = "templates/spec_umovi_vstupnik_template.docx"
            self.fill_and_save_template_dialog(template_path, context, "Журнал спец умов вступників успішно збережений!")
            self.close_db()

        except Exception as e:
            self.show_error_message(f"Помилка при створенні журналу спец умов: {str(e)}")
            print(f"Помилка при створенні журналу спец умов: {str(e)}")
            self.close_db()



    def print_spec_zvit_vstupnik_page(self, study_form, dialog):
        """Друк звіту вступників із динамічним визначенням спеціальностей, типу фінансування та оцінок."""
        try:
            self.connect_db()

            # Визначаємо таблиці залежно від форми навчання
            if study_form == "денна":
                personal_data_table = "applicant_personal_data_day"
                case_table = "personal_case_day"
                benefits_table = "applicant_benefits_day"
                specialties_table = "specialities_day"
            elif study_form == "денна (скорочена)":
                personal_data_table = "applicant_personal_data_day"
                case_table = "personal_case_day_scor"
                benefits_table = "applicant_benefits_day"
                specialties_table = "specialities_day"
            elif study_form == "заочна":
                personal_data_table = "applicant_personal_data_evening"
                case_table = "personal_case_evening"
                benefits_table = "applicant_benefits_evening"
                specialties_table = "specialities_evening"
            else:
                self.show_error_message("Невідома форма навчання.")
                self.close_db()
                return

            # 1. Завантажуємо список спеціальностей динамічно
            self.cursor.execute(f"""
                SELECT kod_specialnosti, name_specialnosti 
                FROM public.{specialties_table}
                ORDER BY kod_specialnosti
            """)
            specialties_list = self.cursor.fetchall()

            if not specialties_list:
                self.show_error_message(f"Немає спеціальностей для форми навчання '{study_form}'.")
                self.close_db()
                return

            # Створюємо словник: name_specialnosti -> kod_specialnosti
            spec_name_to_code = {name: code for code, name in specialties_list}
            # Упорядкований список кодів для стовпців
            spec_codes = [code for code, name in specialties_list]

            # 2. SQL-запит для вибраної форми навчання
            query = f"""
                SELECT 
                    pd.last_name || ' ' || pd.first_name || ' ' || COALESCE(pd.middle_name, '') AS full_name,
                    pd.cert_number,
                    STRING_AGG(DISTINCT pv.kod_pilgi::TEXT, ', ') AS kod_pilgi,
                    os.name_specialnosti,
                    os.number_sprava,
                    SUBSTRING(os.finanse FROM 1 FOR 1) AS finanse,
                    COALESCE(pd.algebra::TEXT, '') || '-' || COALESCE(pd.geometry::TEXT, '') || '-' || 
                    COALESCE(pd.ukr_language::TEXT, '') || '-' || COALESCE(pd.ukr_literature::TEXT, '') AS ocinki
                FROM 
                    public.{personal_data_table} pd
                LEFT JOIN 
                    public.{case_table} os ON pd.cert_number = os.cert_number
                LEFT JOIN 
                    public.{benefits_table} pv ON pd.cert_number = pv.cert_number
                WHERE 
                    os.number_sprava IS NOT NULL
                GROUP BY 
                    pd.cert_number, pd.last_name, pd.first_name, pd.middle_name, 
                    os.name_specialnosti, os.number_sprava, os.finanse, 
                    pd.algebra, pd.geometry, pd.ukr_language, pd.ukr_literature
                ORDER BY 
                    full_name ASC;
            """

            self.cursor.execute(query)
            raw_data = self.cursor.fetchall()

            if not raw_data:
                self.show_error_message(f"Немає даних для звіту за формою навчання '{study_form}'.")
                self.close_db()
                return

            # 3. Організовуємо дані — динамічне групування за спеціальностями
            applicants_dict = {}
            for full_name, cert_number, kod_pilgi, specialty, number_sprava, finanse, ocinki in raw_data:
                if cert_number not in applicants_dict:
                    applicants_dict[cert_number] = {
                        "full_name": full_name,
                        "kod_pilgi": kod_pilgi or "",
                        "specialties": {code: [] for code in spec_codes},
                        "ocinki": ocinki or "—"
                    }

                # Динамічне зіставлення спеціальності із кодом
                if specialty and specialty in spec_name_to_code:
                    code = spec_name_to_code[specialty]
                    applicants_dict[cert_number]["specialties"][code].append(f"{number_sprava} ({finanse})")

            # 4. Формуємо документ через python-docx для підтримки динамічних стовпців
            doc = Document()
            
            # Налаштування альбомної орієнтації та вузьких полів
            section = doc.sections[0]
            new_width, new_height = section.page_height, section.page_width
            section.orientation = 1
            section.page_width = new_width
            section.page_height = new_height
            
            # Вузькі поля сторінки (1.27 см)
            section.top_margin = Cm(1.27)
            section.bottom_margin = Cm(1.27)
            section.left_margin = Cm(1.27)
            section.right_margin = Cm(1.27)
            
            def set_font(run, font_name="Times New Roman", size=12, bold=False):
                run.font.name = font_name
                run.font.size = Pt(size)
                run.bold = bold
                rpr = run._element.get_or_add_rPr()
                rFonts = rpr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    rpr.insert(0, rFonts)
                rFonts.set(qn('w:ascii'), font_name)
                rFonts.set(qn('w:hAnsi'), font_name)
                rFonts.set(qn('w:eastAsia'), font_name)
                rFonts.set(qn('w:cs'), font_name)

            num_cols = 3 + len(spec_codes) + 1
            table = doc.add_table(rows=1, cols=num_cols)
            table.style = 'Table Grid'
            table.autofit = False 
            
            # Заголовки (повторюються на кожній сторінці як колонтитул таблиці)
            hdr_cells = table.rows[0].cells
            headers = ["№ п/п", "ПІБ", "Код пільги"] + spec_codes + ["Оцінки"]
            for i, h_text in enumerate(headers):
                hdr_cells[i].text = str(h_text)
                set_font(hdr_cells[i].paragraphs[0].runs[0], bold=True)
                hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            tr = table.rows[0]._tr
            trPr = tr.get_or_add_trPr()
            tblHeader = OxmlElement('w:tblHeader')
            tblHeader.set(qn('w:val'), "true")
            trPr.append(tblHeader)

            # Дані
            sorted_cert_nums = sorted(applicants_dict.keys(), key=lambda k: applicants_dict[k]["full_name"])
            for i, cert_number in enumerate(sorted_cert_nums, start=1):
                app = applicants_dict[cert_number]
                row_cells = table.add_row().cells
                
                spec_vals = [", ".join(app["specialties"][code]) for code in spec_codes]
                row_data = [str(i), app["full_name"], app["kod_pilgi"]] + spec_vals + [app["ocinki"]]
                
                for j, text in enumerate(row_data):
                    row_cells[j].text = text
                    set_font(row_cells[j].paragraphs[0].runs[0])
                    if j != 1:  # вирівнюємо всі стовпці по центру, крім ПІБ
                        row_cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Після заповнення таблиці встановлюємо точну ширину стовпців
            # Ширина № п/п = 1.5 см, ПІБ = 7.5 см.
            # Залишок рівномірно ділимо між іншими стовпцями.
            total_width = section.page_width - section.left_margin - section.right_margin
            fixed_width = Cm(1.5) + Cm(7.5)
            remaining_width = total_width - fixed_width
            remaining_cols = num_cols - 2
            col_widths = [Cm(1.5), Cm(7.5)] + [remaining_width / remaining_cols] * remaining_cols

            for row in table.rows:
                for j, cell in enumerate(row.cells):
                    cell.width = col_widths[j]

            # 5. Діалог збереження
            root = Tk()
            root.withdraw()
            save_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Documents", "*.docx")],
                initialfile="Звіт_пільги_оцінки",
                title="Оберіть шлях для збереження документа"
            )
            root.destroy()

            if save_path:
                doc.save(save_path)
                dialog.accept()
                self.show_success_message("Звіт вступників успішно збережений!")
            else:
                self.show_error_message("Збереження скасовано користувачем.")

            self.close_db()

        except Exception as e:
            self.show_error_message(f"Помилка при створенні звіту: {str(e)}")
            self.close_db()

    # def print_dodatki_protokol_denne_page(self, start_date, end_date, dialog):
    #     """Друк додатків для протоколів денної і заочної форми навчання."""
    #     try:
    #         self.connect_db()

    #         # SQL-запит для денного та заочного набору даних
    #         query = """
    #         SELECT 
    #             pd.last_name || ' ' || pd.first_name || ' ' || COALESCE(pd.middle_name, '') AS full_name,
    #             sd.kod_specialnosti,
    #             os.finanse,
    #             NULL AS zno_nmt_checkbox,
    #             os.date_sprava AS iso_date
    #         FROM 
    #             public.applicant_personal_data_day pd
    #         JOIN 
    #             public.personal_case_day os ON pd.cert_number = os.cert_number
    #         JOIN 
    #             public.specialities_day sd ON sd.name_specialnosti = os.name_specialnosti
    #         WHERE 
    #             os.date_sprava BETWEEN %s AND %s

    #         UNION ALL

    #         SELECT 
    #             pd.last_name || ' ' || pd.first_name || ' ' || COALESCE(pd.middle_name, '') AS full_name,
    #             sd.kod_specialnosti,
    #             os.finanse,
    #             os.zno_nmt_checkbox,
    #             os.date_sprava AS iso_date
    #         FROM 
    #             public.applicant_personal_data_evening pd
    #         JOIN 
    #             public.personal_case_evening os ON pd.cert_number = os.cert_number
    #         JOIN 
    #             public.specialities_evening sd ON sd.name_specialnosti = os.name_specialnosti
    #         WHERE 
    #             os.date_sprava BETWEEN %s AND %s

    #         ORDER BY 
    #             kod_specialnosti ASC, 
    #             iso_date ASC;
    #         """

    #         # Виконання запиту
    #         self.cursor.execute(query, (start_date, end_date, start_date, end_date))
    #         raw_data = self.cursor.fetchall()

    #         # Перевірка наявності даних
    #         if not raw_data:  # Якщо список порожній
    #             self.show_error_message("У вибраний проміжок дат немає поданих заяв.")
    #             self.close_db()
    #             return  # Завершити функцію

    #         # Організовуємо дані за кодом спеціальності та типом фінансування
    #         specialnosti_data = {}
    #         for full_name, kod_specialnosti, finanse, zno_nmt_checkbox, _ in raw_data:
    #             # Ініціалізація словника для спеціальності
    #             if kod_specialnosti not in specialnosti_data:
    #                 specialnosti_data[kod_specialnosti] = {
    #                     "all": [], "state": [], "paid": [], 
    #                     "zaoch_all": [], "zaoch_state": [], "zaoch_paid": []
    #                 }

    #             # Додавання студента до загального списку
    #             if zno_nmt_checkbox is None:  # Денна форма
    #                 specialnosti_data[kod_specialnosti]["all"].append(full_name)
    #                 if finanse.strip().lower() == "бюджет":
    #                     specialnosti_data[kod_specialnosti]["state"].append(full_name)
    #                 elif finanse.strip().lower() == "контракт":
    #                     specialnosti_data[kod_specialnosti]["paid"].append(full_name)
    #             else:  # Заочна форма
    #                 specialnosti_data[kod_specialnosti]["zaoch_all"].append(full_name)
    #                 if finanse.strip().lower() == "бюджет" and zno_nmt_checkbox == '0':
    #                     specialnosti_data[kod_specialnosti]["zaoch_state"].append(full_name)
    #                 elif finanse.strip().lower() == "контракт":
    #                     specialnosti_data[kod_specialnosti]["zaoch_paid"].append(full_name)

    #         # Формуємо контекст для шаблону
    #         context = {"student": {}}
            
    #         for kod_specialnosti, data in specialnosti_data.items():
    #             # Додавання для денної форми
    #             context["student"][f"all_{kod_specialnosti}_full_name"] = [
    #                 {"index": i + 1, "full_name": name} for i, name in enumerate(data["all"])
    #             ]
    #             context["student"][f"state_{kod_specialnosti}_full_name"] = [
    #                 {"index": i + 1, "full_name": name} for i, name in enumerate(data["state"])
    #             ]
    #             context["student"][f"paid_{kod_specialnosti}_full_name"] = [
    #                 {"index": i + 1, "full_name": name} for i, name in enumerate(data["paid"])
    #             ]
                
    #             # Додавання для заочної форми
    #             context["student"][f"zaoch_all_{kod_specialnosti}_full_name"] = [
    #                 {"index": i + 1, "full_name": name} for i, name in enumerate(data["zaoch_all"])
    #             ]
    #             context["student"][f"zaoch_state_{kod_specialnosti}_full_name"] = [
    #                 {"index": i + 1, "full_name": name} for i, name in enumerate(data["zaoch_state"])
    #             ]
    #             context["student"][f"zaoch_paid_{kod_specialnosti}_full_name"] = [
    #                 {"index": i + 1, "full_name": name} for i, name in enumerate(data["zaoch_paid"])
    #             ]

    #         # Шлях до шаблону
    #         template_path = "templates/dodatki_protokol_template.docx"

    #         # Заповнення шаблону
    #         self.fill_and_save_template(template_path, context, dialog, "Додатки до протоколів успішно збережені!")
    #         self.close_db()

    #     except Exception as e:
    #         self.show_error_message(f"Помилка при створенні звіту: {str(e)}")
    #         self.close_db()
    def print_dodatki_protokol_denne_page(self, start_date, end_date, dialog):
        """Друк додатків для протоколів (універсальний): денна, скорочена та заочна форми."""
        try:
            self.connect_db()

            # SQL-запит з отриманням форми навчання для розрізнення 11 кл
            query = """
            SELECT 
                pd.last_name || ' ' || pd.first_name || ' ' || COALESCE(pd.middle_name, '') AS full_name,
                sd.kod_specialnosti,
                sd.name_specialnosti,
                os.finanse,
                'false' AS zno_nmt, -- Для 9 класу НМТ не існує
                os.date_sprava AS iso_date,
                os.number_sprava,
                '9' AS grade_level,
                '(денна)' AS form_label
            FROM 
                public.applicant_personal_data_day pd
            JOIN 
                public.personal_case_day os ON pd.cert_number = os.cert_number
            JOIN 
                public.specialities_day sd ON sd.name_specialnosti = os.name_specialnosti
            WHERE 
                os.date_sprava BETWEEN %s AND %s

            UNION ALL

            SELECT 
                pd.last_name || ' ' || pd.first_name || ' ' || COALESCE(pd.middle_name, '') AS full_name,
                sd.kod_specialnosti,
                sd.name_specialnosti,
                os.finanse,
                COALESCE(os.zno_nmt_checkbox, 'false') AS zno_nmt,
                os.date_sprava AS iso_date,
                os.number_sprava,
                '11' AS grade_level,
                '(денна)' AS form_label
            FROM 
                public.applicant_personal_data_day pd
            JOIN 
                public.personal_case_day_scor os ON pd.cert_number = os.cert_number
            JOIN 
                public.specialities_day sd ON sd.name_specialnosti = os.name_specialnosti
            WHERE 
                os.date_sprava BETWEEN %s AND %s

            UNION ALL

            SELECT 
                pd.last_name || ' ' || pd.first_name || ' ' || COALESCE(pd.middle_name, '') AS full_name,
                sd.kod_specialnosti,
                sd.name_specialnosti,
                os.finanse,
                COALESCE(os.zno_nmt_checkbox, 'false') AS zno_nmt,
                os.date_sprava AS iso_date,
                os.number_sprava,
                '11' AS grade_level,
                '(заочна)' AS form_label
            FROM 
                public.applicant_personal_data_evening pd
            JOIN 
                public.personal_case_evening os ON pd.cert_number = os.cert_number
            JOIN 
                public.specialities_evening sd ON sd.name_specialnosti = os.name_specialnosti
            WHERE 
                os.date_sprava BETWEEN %s AND %s

            ORDER BY 
                kod_specialnosti ASC, 
                iso_date ASC, 
                number_sprava ASC;
            """

            self.cursor.execute(query, (start_date, end_date, start_date, end_date, start_date, end_date))
            raw_data = self.cursor.fetchall()

            if not raw_data:
                self.show_error_message("У вибраний проміжок дат немає поданих заяв.")
                self.close_db()
                return

            # Структура для групування
            # Ключ: (kod, name_full, form_label), Значення: список студентів
            data_9_all = {}
            data_11_all = {}
            data_11_no_nmt = {}

            for full_name, kod, name, finanse, zno_nmt, iso_date, number_sprava, level, form_label in raw_data:
                student_item = {"full_name": full_name, "iso_date": iso_date, "number_sprava": number_sprava}

                if level == '9':
                    # Для 9 класу зазвичай не пишемо (денна), бо вона там одна, але можемо додати за потреби
                    spec_key = (kod, name, "") 
                    if spec_key not in data_9_all: data_9_all[spec_key] = []
                    data_9_all[spec_key].append(student_item)
                else: # level 11
                    spec_key = (kod, name, form_label)
                    # Всі (Додаток 1 та 3)
                    if spec_key not in data_11_all: data_11_all[spec_key] = []
                    data_11_all[spec_key].append(student_item)
                    
                    # Без НМТ (Додаток 2)
                    if zno_nmt.lower() != 'true':
                        if spec_key not in data_11_no_nmt: data_11_no_nmt[spec_key] = []
                        data_11_no_nmt[spec_key].append(student_item)

            def format_spec_list(grouped_dict):
                """Перетворює словник груп у відсортований список для шаблону."""
                result = []
                # Сортуємо спеціальності за кодом, а потім за формою
                sorted_keys = sorted(grouped_dict.keys(), key=lambda x: (x[0], x[2]))
                for kod, name, form_label in sorted_keys:
                    students = grouped_dict[(kod, name, form_label)]
                    students.sort(key=lambda x: (x["iso_date"], x["number_sprava"]))
                    result.append({
                        "kod": kod,
                        "name": name,
                        "form": form_label,
                        "students": [{"index": i + 1, "full_name": s["full_name"]} for i, s in enumerate(students)]
                    })
                return result

            # Формуємо фінальний контекст
            annex_9_all = format_spec_list(data_9_all)
            annex_11_all = format_spec_list(data_11_all)
            annex_11_no_nmt = format_spec_list(data_11_no_nmt)
            
            common = self._get_common_data()
            short_name = common.get("institution_short_name", "")

            context = {
                "annex_9_all": annex_9_all,
                "annex_11_all": annex_11_all,
                "annex_11_no_nmt": annex_11_no_nmt,
                "short_name": short_name,
                "college_short_name": short_name,
                "student": {
                    "annex_9_all": annex_9_all,
                    "annex_11_all": annex_11_all,
                    "annex_11_no_nmt": annex_11_no_nmt
                }
            }

            template_path = "templates/dodatki_protokol_template.docx"
            self.fill_and_save_template(template_path, context, dialog, "Додатки до протоколів успішно збережені!")
            self.close_db()

        except Exception as e:
            self.show_error_message(f"Помилка при створенні звіту: {str(e)}")
            self.close_db()

    def print_cancelled_cases_journal(self):
        """Формує журнал скасованих справ у Word-документі."""
        try:
            self.connect_db()

            query = """
            SELECT 
                os.number_sprava,
                os.name_specialnosti,
                pd.last_name || ' ' || pd.first_name || ' ' || COALESCE(pd.middle_name, '') AS full_name,
                os.finanse
            FROM public.personal_case_day os
            LEFT JOIN public.applicant_personal_data_day pd ON os.cert_number = pd.cert_number
            WHERE os.is_cancelled IS TRUE

            UNION ALL

            SELECT 
                os.number_sprava,
                os.name_specialnosti,
                pd.last_name || ' ' || pd.first_name || ' ' || COALESCE(pd.middle_name, ''),
                os.finanse
            FROM public.personal_case_day_scor os
            LEFT JOIN public.applicant_personal_data_day pd ON os.cert_number = pd.cert_number
            WHERE os.is_cancelled IS TRUE

            UNION ALL

            SELECT 
                os.number_sprava,
                os.name_specialnosti,
                pd.last_name || ' ' || pd.first_name || ' ' || COALESCE(pd.middle_name, ''),
                os.finanse
            FROM public.personal_case_evening os
            LEFT JOIN public.applicant_personal_data_evening pd ON os.cert_number = pd.cert_number
            WHERE os.is_cancelled IS TRUE

            ORDER BY 1 ASC
            """

            self.cursor.execute(query)
            cancelled = self.cursor.fetchall()

            if not cancelled:
                self.show_error_message("Скасованих справ не знайдено.")
                self.close_db()
                return

            items = [
                {
                    "index": i + 1,
                    "number_sprava": row[0] or "",
                    "name_specialnosti": row[1] or "",
                    "full_name": row[2] or "",
                    "finanse": row[3] or ""
                }
                for i, row in enumerate(cancelled)
            ]

            context = {
                **self._get_common_data(),
                "date": datetime.now().strftime("%d.%m.%Y"),
                "items": items
            }

            template_path = "templates/cancelled_cases_template.docx"
            self.fill_and_save_template_dialog(template_path, context, "Журнал скасованих справ збережено!")
            self.close_db()

        except Exception as e:
            self.show_error_message(f"Помилка при створенні журналу: {str(e)}")
            self.close_db()

    def print_export_date_vstupnik(self):

        """Експортує дані з бази в Excel з поділом на денну, денну (скорочену) та заочну форми навчання."""
        try:
            self.connect_db()

            # Заголовки українською мовою
            headers = [
                "Номер справи", "Назва спеціальності", "Дата подачі справи", "Фінансування", "Скасована заява",
                "Ім'я", "Прізвище", "По батькові", "ПІП", "Телефон", "Громадянство",
                "Номер паспорта", "Ким виданий", "Дата видачі паспорта", "Ідентифікаційний код", 
                "Свідоцтво про освіту", "Дата видачі свідоцтва", "Адреса",
                "Ім'я батька", "Прізвище батька", "По батькові батька", "Робота батька", "Телефон батька",
                "Ім'я матері", "Прізвище матері", "По батькові матері", "Телефон матері", "Робота матері",
                "Потреба в гуртожитку", "Стать", "Алгебра", "Геометрія", "Українська мова", "Українська література",
                "Назва школи", "Дата народження", "Коди пільг", "Документи про пільги"
            ]

            # Запит для денної форми
            query_denna = """
            SELECT 
                os.number_sprava,
                os.name_specialnosti,
                os.date_sprava,
                os.finanse,
                os.is_cancelled,
                pd.first_name,
                pd.last_name,
                pd.middle_name,
                pd.pip,
                pd.phone,
                pd.citizenship,
                pd.passport_number,
                pd.issued_by,
                pd.issue_date,
                pd.id_code,
                pd.cert_number,
                pd.cert_issue_date,
                pd.address,
                pd.father_first_name,
                pd.father_last_name,
                pd.father_middle_name,
                pd.father_job,
                pd.father_phone,
                pd.mother_first_name,
                pd.mother_last_name,
                pd.mother_middle_name,
                pd.mother_phone,
                pd.mother_job,
                pd.hostel_need,
                pd.gender,
                pd.algebra,
                pd.geometry,
                pd.ukr_language,
                pd.ukr_literature,
                pd.school_name,
                pd.date_birth,
                STRING_AGG(pv.kod_pilgi::TEXT, ', ') AS kod_pilgi,
                STRING_AGG(pv.document_pilgi, ', ') AS document_pilgi
            FROM 
                public.personal_case_day os
            LEFT JOIN 
                public.applicant_personal_data_day pd ON os.cert_number = pd.cert_number
            LEFT JOIN 
                public.applicant_benefits_day pv ON os.cert_number = pv.cert_number
            GROUP BY 
                os.number_sprava, os.name_specialnosti, os.date_sprava, os.finanse, os.is_cancelled,
                pd.first_name, pd.last_name, pd.middle_name, pd.pip, pd.phone, pd.citizenship, 
                pd.passport_number, pd.issued_by, pd.issue_date, pd.id_code, pd.cert_number, 
                pd.cert_issue_date, pd.address, pd.father_first_name, pd.father_last_name, 
                pd.father_middle_name, pd.father_job, pd.father_phone, pd.mother_first_name, 
                pd.mother_last_name, pd.mother_middle_name, pd.mother_phone, pd.mother_job, 
                pd.hostel_need, pd.gender, pd.algebra, pd.geometry, pd.ukr_language, 
                pd.ukr_literature, pd.school_name, pd.date_birth
            ORDER BY 
                os.number_sprava ASC;
            """

            self.cursor.execute(query_denna)
            data_denna = self.cursor.fetchall()

            # Запит для денної (скороченої) форми
            query_denna_scor = query_denna.replace(
                "public.personal_case_day", "public.personal_case_day_scor"
            )

            self.cursor.execute(query_denna_scor)
            data_denna_scor = self.cursor.fetchall()

            # Запит для заочної форми
            query_zaoch = query_denna.replace(
                "public.personal_case_day", "public.personal_case_evening"
            ).replace(
                "public.applicant_personal_data_day", "public.applicant_personal_data_evening"
            ).replace(
                "public.applicant_benefits_day", "public.applicant_benefits_evening"
            )

            self.cursor.execute(query_zaoch)
            data_zaoch = self.cursor.fetchall()

            # Виклик діалогу для збереження файлу через Tkinter
            root = Tk()
            root.withdraw()  # Приховуємо головне вікно Tkinter

            file_path = filedialog.asksaveasfilename(
                title="Зберегти файл як...",
                defaultextension=".xlsx",
                filetypes=[("Excel Files", "*.xlsx"), ("All Files", "*.*")]
            )

            if not file_path:
                self.show_success_message("Експорт скасовано користувачем.")
                self.close_db()
                return

            # Створення Excel-файлу
            from openpyxl.styles import PatternFill
            workbook = Workbook()
            red_fill = PatternFill(start_color="FFB4B4", end_color="FFB4B4", fill_type="solid")

            def append_rows_with_highlight(sheet, data):
                """Додає рядки на лист, виділяючи скасовані справи червоним."""
                for row in data:
                    # is_cancelled - 5-й стовпець (index 4)
                    is_cancelled = row[4]
                    # Перетворюємо bool -> 'Так'/'Ні' для читабельності в Excel
                    row_to_write = list(row)
                    row_to_write[4] = "Так" if is_cancelled else "Ні"
                    sheet.append(row_to_write)
                    if is_cancelled:
                        for cell in sheet[sheet.max_row]:
                            cell.fill = red_fill

            # Лист для денної форми
            sheet_denna = workbook.active
            sheet_denna.title = "Денна форма"
            sheet_denna.append(headers)
            append_rows_with_highlight(sheet_denna, data_denna)

            # Лист для денної (скороченої) форми
            sheet_denna_scor = workbook.create_sheet(title="Денна (скорочена) форма")
            sheet_denna_scor.append(headers)
            append_rows_with_highlight(sheet_denna_scor, data_denna_scor)

            # Лист для заочної форми
            sheet_zaoch = workbook.create_sheet(title="Заочна форма")
            sheet_zaoch.append(headers)
            append_rows_with_highlight(sheet_zaoch, data_zaoch)

            # Збереження Excel-файлу
            workbook.save(file_path)

            self.show_success_message(f"Дані успішно експортовано в файл")
            self.close_db()

        except Exception as e:
            self.show_error_message(f"Помилка при експорті даних: {str(e)}")
            self.close_db()

   
    def fill_and_save_template_dialog(self, template_path, data, success_message):
        """Заповнює шаблон DOCX і зберігає документ із вибором місця збереження."""
        try:
            common = self._get_common_data()
            common.update(data)
            doc = DocxTemplate(template_path)
            doc.render(common)

            # Показуємо діалог для вибору місця збереження
            root = Tk()
            root.withdraw()
            save_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Documents", "*.docx")],
                initialfile="Звіт",
                title="Оберіть шлях для збереження документа"
            )
            root.destroy()

            if save_path:
                doc.save(save_path)
                self.show_success_message(success_message)
            else:
                self.show_error_message("Збереження скасовано користувачем.")

        except Exception as e:
            self.show_error_message(f"Помилка при створенні документів: {str(e)}")
            print(f"Помилка при створенні документів: {str(e)}")

    def fill_and_save_template(self, template_path, data, dialog, success_message, colors_config=None):
        """Заповнює шаблон DOCX та зберігає його з можливістю розфарбовування рядків таблиці."""
        try:
            common = self._get_common_data()
            common.update(data)
            doc = DocxTemplate(template_path)
            doc.render(common)

            # Якщо передано конфігурацію кольорів, обробляємо таблиці
            if colors_config:
                # doc.docx - це об'єкт Document всередині DocxTemplate
                if doc.tables:
                    table = doc.tables[0]
                    for row_idx, color in colors_config.items():
                        # Перевіряємо, чи існує такий рядок (row_idx 0 - заголовок, дані з 1)
                        if 0 <= row_idx < len(table.rows):
                            self._set_row_color(table.rows[row_idx], color)

            # Показуємо діалог для вибору місця збереження
            root = Tk()
            root.withdraw()
            save_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Documents", "*.docx")],
                initialfile="Рейтинговий_список",
                title="Оберіть шлях для збереження документа"
            )
            root.destroy()

            if save_path:
                doc.save(save_path)
                if dialog: dialog.accept()
                self.show_success_message(success_message)
            else:
                self.show_error_message("Збереження було скасовано.")

        except Exception as e:
            self.show_error_message(f"Помилка при створенні документа: {str(e)}")
            print(f"Помилка при створенні документа: {str(e)}")


